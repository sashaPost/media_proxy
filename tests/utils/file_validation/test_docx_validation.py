import docx.opc.constants
import pytest
from unittest.mock import patch, MagicMock
from utils.file_validation import is_valid_docx


class TestIsValidDOCX:
    @pytest.fixture
    def mock_zipfile(self):
        with patch("utils.file_validation.zipfile.ZipFile") as mock:
            yield mock

    @pytest.fixture
    def mock_docx_document(self):
        with patch("utils.file_validation.docx.Document") as mock:
            yield mock

    def test_valid_docx(self, mock_zipfile, mock_docx_document):
        mock_file = MagicMock()
        mock_file.read.return_value = b"PK\x03\x04"

        # Mock zipfile.is_zipfile to return True
        with patch(
            "utils.file_validation.zipfile.is_zipfile", return_value=True
        ), patch("utils.file_validation.io.BytesIO") as mock_bytesio:

            mock_bytesio_instance = MagicMock()
            mock_bytesio.return_value = mock_bytesio_instance

            mock_zipfile.return_value.namelist.return_value = [
                "[Content_Types].xml",
                "word/document.xml",
            ]

            mock_doc = MagicMock()
            mock_doc.paragraphs = []
            mock_doc.part.rels.values.return_value = []  # No external links

            mock_docx_document.return_value = mock_doc

            result = is_valid_docx(mock_file)
            assert result is True

            # Verify that ZipFile was called with the BytesIO instance
            mock_zipfile.assert_called_once_with(mock_bytesio_instance)

            # Verify that docx.Document was called
            mock_docx_document.assert_called()

    def test_invalid_zip(self, mock_zipfile):
        mock_file = MagicMock()
        mock_file.read.return_value = b"Not a ZIP file"

        mock_zipfile.is_zipfile.return_value = False

        assert is_valid_docx(mock_file) is False

    def test_docx_with_macros(self, mock_zipfile, mock_docx_document):
        mock_file = MagicMock()
        mock_file.read.return_value = b"PK\x03\x04"

        mock_zipfile.return_value.namelist.return_value = ["word/vbaProject.bin"]

        assert is_valid_docx(mock_file) is False

    def test_docx_with_external_links(self, mock_zipfile, mock_docx_document):
        mock_file = MagicMock()
        mock_file.read.return_value = b"PK\x03\x04"

        mock_zipfile.return_value.namelist.return_value = []

        mock_doc = MagicMock()

        mock_rel = MagicMock()
        mock_rel.reltype = docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK

        mock_doc.part.rels.values.return_value = [mock_rel]

        mock_docx_document.return_value = mock_doc

        assert is_valid_docx(mock_file) is False

    def test_docx_with_embedded_objects(self, mock_zipfile, mock_docx_document):
        mock_file = MagicMock()
        mock_file.read.return_value = b"PK\x03\x04"

        mock_zipfile.return_value.namelist.return_value = [
            "word/embeddings/oleObject1.bin"
        ]

        assert is_valid_docx(mock_file) is False

    def test_docx_with_script_injection(self, mock_zipfile, mock_docx_document):
        mock_file = MagicMock()
        mock_file.read.return_value = b"PK\x03\x04"

        mock_zipfile.return_value.namelist.return_value = []

        mock_doc = MagicMock()
        mock_doc.paragraphs = [MagicMock()]
        mock_doc.paragraphs[0].text = '<script>alert("XSS")</script>'

        mock_docx_document.return_value = mock_doc

        assert is_valid_docx(mock_file) is False

    def test_docx_with_suspicious_keywords(self, mock_zipfile, mock_docx_document):
        mock_file = MagicMock()
        mock_file.read.return_value = b"PK\x03\x04"

        mock_zipfile.return_value.namelist.return_value = []

        mock_doc = MagicMock()
        mock_doc.paragraphs = [MagicMock()]
        mock_doc.paragraphs[0].text = 'exec("malicious code")'

        mock_docx_document.return_value = mock_doc

        assert is_valid_docx(mock_file) is False

    def test_large_docx(self, mock_zipfile, mock_docx_document):
        mock_file = MagicMock()
        mock_file.read.return_value = b"PK\x03\x04" * (5 * 1024 * 1024)

        mock_zipfile.return_value.namelist.return_value = []

        assert is_valid_docx(mock_file) is False
